from flask import Flask, request, render_template
import docx
import PyPDF2  # Import PyPDF2
from io import BytesIO

app = Flask(__name__)


def extract_metadata_from_docx(doc):
    metadata = {}
    core_props = doc.core_properties
    metadata['title'] = getattr(core_props, 'title', None)
    metadata['subject'] = getattr(core_props, 'subject', None)
    metadata['creator'] = getattr(core_props, 'creator', None)
    metadata['keywords'] = getattr(core_props, 'keywords', None)
    metadata['last_modified_by'] = getattr(core_props, 'last_modified_by', None)
    metadata['revision'] = getattr(core_props, 'revision', None)
    metadata['modified'] = getattr(core_props, 'modified', None)
    metadata['created'] = getattr(core_props, 'created', None)
    metadata['paragraphs_count'] = len(doc.paragraphs)
    metadata['tables_count'] = len(doc.tables)
    return metadata


# Modified function to use PdfReader and metadata attribute
def extract_metadata_from_pdf(file):
    metadata = {}
    pdf_reader = PyPDF2.PdfReader(file)
    info = pdf_reader.metadata  # Use metadata attribute instead of getDocumentInfo()
    metadata['title'] = getattr(info, 'title', None)
    metadata['author'] = getattr(info, 'author', None)
    metadata['subject'] = getattr(info, 'subject', None)
    metadata['creator'] = getattr(info, 'creator', None)
    metadata['producer'] = getattr(info, 'producer', None)
    metadata['creation_date'] = getattr(info, 'creation_date', None)
    metadata['modification_date'] = getattr(info, 'modification_date', None)
    metadata['num_pages'] = len(pdf_reader.pages)
    return metadata


@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('upload.html', error='No file part')
        file = request.files['file']

        if file.filename == '':
            return render_template('upload.html', error='No selected file')

        if not file.filename.endswith(('.docx', '.pdf')):
            return render_template('upload.html', error='Invalid file type. Please upload a .docx or .pdf file.')

        file_bytes = BytesIO(file.read())

        if file.filename.endswith('.docx'):
            try:
                doc = docx.Document(file_bytes)
                metadata = extract_metadata_from_docx(doc)
            except Exception as e:
                print("Error reading docx:", str(e))
                return render_template('upload.html', error='Error reading the .docx file. Please try again.')
        else:  # .pdf
            try:
                file_bytes.seek(0)  # Reset to the beginning of the file
                metadata = extract_metadata_from_pdf(file_bytes)
            except Exception as e:
                print("Error reading pdf:", str(e))
                return render_template('upload.html', error='Error reading the .pdf file. Please try again.')

        return render_template('metadata.html', metadata=metadata)

    return render_template('upload.html')


if __name__ == "__main__":
    app.run(debug=True)
